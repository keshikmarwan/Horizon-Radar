import io
import re
import zipfile
from html import unescape

import pdfplumber
from docx import Document


def _extract_text_from_docx_bytes(data: bytes) -> str:
    # Primary path: python-docx for robust paragraph/table extraction.
    try:
        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for p in doc.paragraphs:
            txt = (p.text or '').strip()
            if txt:
                lines.append(txt)

        for table in doc.tables:
            for row in table.rows:
                row_vals = []
                for cell in row.cells:
                    cell_txt = ' '.join((cell.text or '').split())
                    if cell_txt:
                        row_vals.append(cell_txt)
                if row_vals:
                    lines.append(' | '.join(row_vals))

        joined = '\n'.join(lines).strip()
        if joined:
            return joined
    except Exception:
        pass

    # Fallback: raw XML read in case the parser fails.
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        parts = []
        for name in zf.namelist():
            if name.startswith('word/') and name.endswith('.xml'):
                try:
                    parts.append(zf.read(name).decode('utf-8', errors='ignore'))
                except Exception:
                    continue

    xml_text = '\n'.join(parts)
    xml_text = re.sub(r'</w:p>', '\n', xml_text, flags=re.IGNORECASE)
    xml_text = re.sub(r'</w:tr>', '\n', xml_text, flags=re.IGNORECASE)
    xml_text = re.sub(r'</w:tc>', ' | ', xml_text, flags=re.IGNORECASE)

    text = re.sub(r'<[^>]+>', ' ', xml_text)
    text = unescape(text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    return text.strip()


def _extract_text_from_pdf_bytes(data: bytes) -> str:
    output: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:200]:
            output.append(page.extract_text() or '')
    return '\n'.join(output).strip()


def extract_text_from_upload(filename: str, content_type: str | None, data: bytes) -> tuple[str, str]:
    name = (filename or '').lower()
    ctype = (content_type or '').lower()

    if name.endswith('.docx') or 'officedocument.wordprocessingml.document' in ctype:
        try:
            text = _extract_text_from_docx_bytes(data)
            return text, 'docx'
        except Exception:
            # Hard fallback: never fail the endpoint for extraction.
            pass

    if name.endswith('.pdf') or 'application/pdf' in ctype:
        try:
            return _extract_text_from_pdf_bytes(data), 'pdf'
        except Exception:
            # Hard fallback: never fail the endpoint for extraction.
            pass

    if name.endswith('.doc'):
        try:
            text = data.decode('latin-1', errors='ignore')
            return text, 'doc-binary-decoded'
        except Exception:
            pass

    try:
        text = data.decode('utf-8')
        return text, 'text'
    except UnicodeDecodeError:
        text = data.decode('latin-1', errors='ignore')
        return text, 'binary-decoded'
